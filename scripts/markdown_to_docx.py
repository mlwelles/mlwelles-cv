#!/usr/bin/env python3
"""Convert Markdown document to a simple styled DOCX document.
This is a simpler converter for documents like cover letters."""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_simple_docx(md_path, output_path):
    """Create a simple DOCX from markdown paragraphs."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Add header with name and contact info
    # Name
    name = doc.add_paragraph("Michael Welles")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.runs[0]
    name_run.font.size = Pt(14)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name.paragraph_format.space_after = Pt(2)

    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = "Email: mlwelles@gmail.com | Phone: 347-450-6518 | Location: Brooklyn, NY"
    run = contact.add_run(contact_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(64, 64, 64)
    contact.paragraph_format.space_after = Pt(18)

    # Process each paragraph of the cover letter body
    paragraphs = content.split('\n\n')

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Add paragraph
        p = doc.add_paragraph(para_text)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(output_path)
    print(f"Created styled DOCX: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: markdown_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    create_simple_docx(input_path, output_path)


if __name__ == '__main__':
    main()
