#!/usr/bin/env python3
"""Convert Markdown resume to a well-styled DOCX document."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_markdown_resume(md_path):
    """Parse the Markdown resume into structured sections."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = {}
    current_section = None
    current_subsection = None
    lines = content.split('\n')

    for line in lines:
        # H1 - Name
        if line.startswith('# '):
            sections['name'] = line[2:].strip()

        # Contact info (bullet points at top)
        elif line.startswith('- **') and ':' in line:
            if 'contact' not in sections:
                sections['contact'] = []
            # Extract "**Label:** value"
            match = re.match(r'- \*\*(.+?):\*\* (.+)', line)
            if match:
                sections['contact'].append((match.group(1), match.group(2)))

        # H2 - Major sections
        elif line.startswith('## '):
            current_section = line[3:].strip()
            sections[current_section] = {'type': 'section', 'content': []}
            current_subsection = None

        # H3 - Job titles / subsections
        elif line.startswith('### '):
            current_subsection = {
                'title': line[4:].strip(),
                'meta': '',
                'description': '',
                'bullets': []
            }
            if current_section:
                sections[current_section]['content'].append(current_subsection)

        # Italic line after H3 (dates/location)
        elif line.startswith('*') and line.endswith('*') and current_subsection:
            current_subsection['meta'] = line.strip('*')

        # Bullet points
        elif line.startswith('- ') and current_subsection:
            current_subsection['bullets'].append(line[2:].strip())

        # Regular paragraph text
        elif line.strip() and current_section:
            if current_subsection:
                if not current_subsection['description']:
                    current_subsection['description'] = line.strip()
                else:
                    current_subsection['description'] += ' ' + line.strip()
            else:
                # Section-level content (like Overview paragraph)
                if isinstance(sections[current_section]['content'], list):
                    if not sections[current_section]['content'] or isinstance(sections[current_section]['content'][-1], dict):
                        sections[current_section]['content'].append(line.strip())
                    else:
                        sections[current_section]['content'][-1] += ' ' + line.strip()

        # Bold section headers (like "**Languages & Frameworks:**")
        elif line.startswith('**') and ':**' in line:
            if current_section:
                sections[current_section]['content'].append(line)

    return sections


def create_styled_docx(sections, output_path):
    """Create a professionally styled DOCX document."""
    doc = Document()

    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Name (Title)
    if 'name' in sections:
        name = doc.add_paragraph(sections['name'])
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name.runs[0]
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name.paragraph_format.space_after = Pt(2)

    # Contact info
    if 'contact' in sections:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = ' | '.join([f"{label}: {value}" for label, value in sections['contact']])
        run = contact.add_run(contact_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(64, 64, 64)
        contact.paragraph_format.space_after = Pt(6)

    # Process other sections
    for section_name, section_data in sections.items():
        if section_name in ['name', 'contact']:
            continue

        if section_data.get('type') == 'section':
            # Section header
            heading = doc.add_paragraph(section_name)
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(13)
            heading_run.font.bold = True
            heading_run.font.color.rgb = RGBColor(0, 51, 102)

            # Add bottom border to section header
            heading.paragraph_format.space_before = Pt(4)
            heading.paragraph_format.space_after = Pt(2)

            # Process section content
            for item in section_data['content']:
                if isinstance(item, str):
                    # Plain text or formatted text
                    if item.startswith('**') and ':**' in item:
                        # Bold category line (like "**Languages:**")
                        p = doc.add_paragraph()
                        # Parse "**Category:** content"
                        match = re.match(r'\*\*(.+?):\*\* (.+)', item)
                        if match:
                            bold_run = p.add_run(match.group(1) + ': ')
                            bold_run.font.bold = True
                            bold_run.font.size = Pt(10)

                            content_run = p.add_run(match.group(2))
                            content_run.font.size = Pt(10)
                        else:
                            run = p.add_run(item)
                            run.font.size = Pt(10)
                        p.paragraph_format.space_after = Pt(1)
                    else:
                        # Regular paragraph
                        p = doc.add_paragraph(item)
                        p.runs[0].font.size = Pt(11)
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.0

                elif isinstance(item, dict):
                    # Job/subsection
                    # Job title
                    title = doc.add_paragraph(item['title'])
                    title_run = title.runs[0]
                    title_run.font.size = Pt(12)
                    title_run.font.bold = True
                    title.paragraph_format.space_before = Pt(6)
                    title.paragraph_format.space_after = Pt(1)

                    # Meta (dates/location)
                    if item['meta']:
                        meta = doc.add_paragraph(item['meta'])
                        meta_run = meta.runs[0]
                        meta_run.font.size = Pt(10)
                        meta_run.font.italic = True
                        meta_run.font.color.rgb = RGBColor(96, 96, 96)
                        meta.paragraph_format.space_after = Pt(3)

                    # Description
                    if item['description']:
                        desc = doc.add_paragraph(item['description'])
                        desc.runs[0].font.size = Pt(10)
                        desc.paragraph_format.space_after = Pt(3)
                        desc.paragraph_format.line_spacing = 1.0

                    # Bullets
                    for bullet in item['bullets']:
                        p = doc.add_paragraph(bullet, style='List Bullet')
                        p.runs[0].font.size = Pt(10)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.line_spacing = 1.0

    doc.save(output_path)
    print(f"Created styled DOCX: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: markdown_resume_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    sections = parse_markdown_resume(input_path)
    create_styled_docx(sections, output_path)


if __name__ == '__main__':
    main()
